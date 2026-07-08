"""
数据准备模块（企业级改造）

本模块在整个 RAG 系统中的定位（画一张数据流图）：
    企业文档(PDF/Excel/Word/MD)          ┌──────────────────────┐
         │                               │  ③ 向量索引(本模块下游) │
         ▼                               │  把 chunk 转成向量      │
    ┌──────────────────────┐             │  存进 FAISS            │
    │  ① 数据接入(本模块)   │ ──chunk──▶ │                        │
    │  多格式解析→统一结构  │             └──────────────────────┘
    │  + 企业元数据(部门)   │                       │
    └──────────────────────┘                       ▼
         │                               ┌──────────────────────┐
         ▼                               │  ④ 检索 + ⑤ 生成      │
    ┌──────────────────────┐             │  (后续模块)           │
    │  ② 文档处理(本模块)   │             └──────────────────────┘
    │  表格感知切分         │
    │  + 父子分块保留       │
    └──────────────────────┘

一句话：把"人看的文档"变成"向量检索能用的 chunk"。

改造点（对照《真实RAG系统全貌》①②子系统）：
- 领域：食谱 markdown → 企业文档（PDF/DOCX/MD/TXT）
- 解析：仅 md → 多格式（PyMuPDF/python-docx），带页码元数据
- 元数据：菜品名/难度 → 来源/部门/页码（权限+溯源用）
- 切分：MarkdownHeader → 递归切分（通用，PDF 无 # 结构）+ 父子分块保留
"""

# ===== 标准库导入（标准库优先，第三方按需 import 到方法内以加速冷启动） =====
import logging          # 日志，贯穿全项目；这里只取 logger，不配置（配置在 logging_config.py）
import hashlib          # 算 content_hash 和 parent_id（确定性 ID，便于增量索引对比）
import uuid             # 生成 chunk_id（唯一性，不要求确定性）
from pathlib import Path                # 面向对象的路径操作，跨平台（Windows/ Linux 路径分隔符自动处理）
from typing import List, Dict, Any      # 类型注解，IDE 提示 + 可读性，运行时不强制

# ===== LangChain 文本切分器（文档处理的核心工具） =====
# RecursiveCharacterTextSplitter：递归字符切分器，按分隔符优先级递归切分
#   —— 先尝试 \n\n 切，切不动再用 \n，再不行用 。；， 等中文标点，最后才是单字符
#   优点：尽量沿语义边界切，比固定长度切分质量高
# MarkdownHeaderTextSplitter：按 Markdown 标题(#/##/###)切分（本项目实际未在主流程使用，
#   保留导入是因为部分企业文档是结构化 MD，可作为切分增强手段）
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
# Document：LangChain 的标准文档载体，就俩字段：
#   page_content: str   —— 文本内容（要被 embedding 的东西）
#   metadata: dict      —— 任意元数据（source/page/department 等，检索过滤和引用溯源全靠它）
from langchain_core.documents import Document

# 每个模块用自己的 __name__ 取 logger，这样日志能按模块名过滤
# （比如只想看 data_preparation 的日志：logging.getLogger("rag_modules.data_preparation").setLevel(DEBUG)）
logger = logging.getLogger(__name__)


def _transpose_comparison_table(df):
    """
    转置 GPU 对比表：GPU 型号为列 → GPU 型号为行。

    【为什么需要这一步？—— 这是检索友好性的核心优化】
    企业里很多表格是"对比表"，结构是"属性为行、产品型号为列"（人类阅读习惯）：
    用户提问"GB200 的显存是多少"，但 GB200 在列里，向量检索按"行"召回，
    每一行只是"显存=192GB"，根本没出现"GB200"这个词 → 检索召回不到。
    转置后，每行是一个 GPU 型号的完整描述 → "GB200, Blackwell, 192GB..." 一行命中。

    输入（常见格式，型号为列）：
           | GPU型号 | GB200  | B200   | H100   | ...
           | 产品架构 | Blackwell | Blackwell | Hopper | ...
           | 显存    | 192GB  | 192GB  | 80GB   | ...

    输出（型号为行，每行=一个产品完整画像）：
           | GPU型号 | 产品架构    | 显存   | ...
           | GB200   | Blackwell  | 192GB  | ...
           | B200    | Blackwell  | 192GB  | ...
           | H100    | Hopper     | 80GB   | ...

    这样每行变成一个 GPU 型号的完整描述，用户搜"GB200 显存"能直接命中。
    """
    import pandas as pd  # 局部导入：pandas 重，只在用到时加载，避免拖慢模块导入
    first_col = df.columns[0]  # 第一列（如 "GPU型号"，这是真正的"属性名"列）
    # set_index：把第一列变成行索引；.T：转置（行列互换）
    # 效果：原来的列名(GB200/B200/..)变成了新的第一列，原来的行变成了新列名
    df_t = df.set_index(first_col).T
    df_t.index.name = first_col       # 新的行索引列也命名为 "GPU型号"
    df_t = df_t.reset_index()         # 行索引还原成普通列（这样 GB200/B200 才是数据行，不是索引）
    return df_t


class DataPreparationModule:
    """数据准备模块 - 企业文档多格式解析、切分、元数据

    职责边界（单一职责原则）：
      - 只管"读文档 + 切分"，不管"向量化"（那是 index_construction 的事）
      - 只管"产 chunk"，不管"检索"（那是 retrieval_optimization 的事）
    """

    def __init__(self, data_path: str, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        Args:
            data_path: 文档根目录（支持 PDF/DOCX/MD/TXT）
            chunk_size: 切块大小（字符数）。为什么是 500？
                - 太小（<200）：上下文割裂，"满3年可休15天"可能切成"满3年"+"15天"两块，检索到也答不出
                - 太大（>1500）：向量稀释（embedding 把一整页压成一个向量，关键词信号被平均掉）
                - 500 是中文 RAG 经验值，配合 bge-small-zh（max_seq_length=512）刚好一个向量装得下
            chunk_overlap: 重叠（相邻块共享的字符数）。为什么需要重叠？
                - 避免刚好在关键句中间切断；重叠 25-50 让边界信息在两块都出现
                - 但太大（>100）会有冗余、检索重复；这里 25 是 chunk_size 的 5%，合理
        """
        self.data_path = data_path
        # resolve() 把相对路径转绝对路径，缓存到 _data_root
        # 【性能】避免后续每个文件都调一次 resolve（IO 操作），启动时算一次复用
        # 【作用】_relative_source/_parent_id/_enhance_metadata 都依赖这个根路径算"相对路径"
        self._data_root = Path(data_path).resolve()
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        # 三个核心数据结构，承载整个模块的产物：
        # documents：父文档（一篇 PDF = 一个 Document，page_content 是整篇文本）
        self.documents: List[Document] = []
        # chunks：子块（把 documents 切分后的结果，向量检索的真正单元）
        self.chunks: List[Document] = []
        # parent_child_map：子块ID → 父文档ID 的映射表
        # 【父子检索】小 chunk 检索精准但上下文不足，命中后用它回溯父文档拿完整上下文
        self.parent_child_map: Dict[str, str] = {}

    # ===== ① 数据接入：多格式加载 =====
    def load_documents(self) -> List[Document]:
        """加载企业文档（PDF/DOCX/MD/TXT），按文件类型分派解析器

        设计要点：
        1. rglob 递归遍历（支持子目录，子目录名=部门=权限边界）
        2. 按后缀分派解析器（策略模式：每种格式一个 _load_xxx 方法）
        3. 单文件失败不中断（异常隔离，保证其余文件正常处理）
        """
        logger.info(f"正在从 {self.data_path} 加载企业文档...")
        documents = []
        data_path_obj = Path(self.data_path)

        # 启动期硬校验：数据目录都不存在，后面全白跑，不如早失败
        if not data_path_obj.exists():
            raise FileNotFoundError(f"数据路径不存在: {self.data_path}")

        # rglob("*")：递归遍历所有层级（glob 只遍历一级，rglob 才会下钻子目录）
        # sorted：保证遍历顺序确定（否则不同 OS/文件系统顺序可能不同，增量索引对比会乱）
        for file_path in sorted(data_path_obj.rglob("*")):
            if file_path.is_dir():
                continue  # 跳过目录本身，只处理文件
            # 【坑】Windows Office（Word/Excel）打开文件时会生成 ~$ 开头的临时锁文件
            # 这些文件是空的或损坏的，解析必然失败，必须跳过（否则一堆无意义告警）
            if file_path.name.startswith("~$"):
                logger.debug(f"跳过临时文件: {file_path.name}")
                continue
            suffix = file_path.suffix.lower()   # 统一小写，兼容 .PDF / .Pdf 等大小写写法
            # try/except 包住单个文件：一个文件炸了不能拖垮整批加载
            # 【生产经验】批量处理必须有异常隔离，否则一个坏文件让整个知识库构建失败
            try:
                # 策略模式：按后缀选解析器。新增格式只需加一个 elif 分支 + 一个 _load_xxx 方法
                if suffix == ".pdf":
                    docs = self._load_pdf(file_path)
                elif suffix == ".docx":
                    docs = self._load_docx(file_path)
                elif suffix in (".xlsx", ".xls"):
                    docs = self._load_xlsx(file_path)
                elif suffix == ".md":
                    docs = self._load_text(file_path, file_type="markdown")
                elif suffix == ".txt":
                    docs = self._load_text(file_path, file_type="text")
                else:
                    logger.debug(f"跳过不支持的格式: {file_path}")
                    continue  # 未知格式直接跳过，不报错（比如 .png 图片）
                documents.extend(docs)
            except Exception as e:
                # 【生产做法】解析失败只记 warning 不中断，真实系统会把这类文件丢进"死信队列"
                # 后续由运维人工介入（补依赖、转格式、或放弃这个文件）
                logger.warning(f"解析失败 {file_path}: {e}")

        # 所有文档加载完后，统一增强元数据（抽取部门/分类）
        # 为什么放循环外而不是每个 _load_xxx 里做？——集中一处，避免漏改、便于维护
        for doc in documents:
            self._enhance_metadata(doc)

        self.documents = documents
        # 汇总日志：列出本次加载了哪些文件，方便人工确认是否漏加载/多加载
        # set 去重：一个 PDF 多页 → 多个 Document，但 source 相同，文件只列一次
        loaded_files = sorted(set(d.metadata.get("source", "?") for d in documents))
        logger.info(f"成功加载 {len(documents)} 个文档片段（{len(loaded_files)} 个文件）:")
        for f in loaded_files:
            logger.info(f"  📄 {f}")
        return documents

    def _load_pdf(self, file_path: Path) -> List[Document]:
        """解析 PDF，保留页码（溯源用）。扫描件无文字层时 OCR 兜底。

        为什么用 PyMuPDF（import fitz）而不是 pdfplumber/PyPDF2？
        - PyMuPDF（fitz）：速度最快，中文支持好，能拿页码、坐标、图片
        - pdfplumber：表格识别强，但慢（10倍于 fitz）
        - PyPDF2：老、慢、中文常乱码
        这里优先速度，表格交给下游的表格感知切分处理。
        """
        import fitz  # PyMuPDF 的导入名是 fitz（历史命名，不是 bug）
        docs = []
        parent_id = self._parent_id(file_path)       # 父文档 ID（同一 PDF 的所有页共享）
        rel_source = self._relative_source(file_path) # 相对路径（如 "mye/服务器知识.pdf"，含部门）
        # with 上下文管理器：保证文件句柄释放（Windows 下不释放会锁文件，导致后续无法移动/删除）
        with fitz.open(file_path) as pdf:
            # enumerate(pdf, 1)：从 1 开始计页码（人类习惯，溯源时显示"第3页"而非"第2页"）
            for page_num, page in enumerate(pdf, 1):
                # get_text("text")：提取纯文本（还有 "html"/"dict" 等模式，"text" 最快）
                text = page.get_text("text")
                # 【兜底机制】如果这页提取不出文字（strip 后为空），说明可能是扫描件（图片型 PDF）
                # 此时文字层是空的，必须走 OCR 把图片里的文字识别出来
                if not text.strip():
                    text = self._ocr_page(page, page_num, file_path)
                # 二次检查：OCR 也可能识别不出（比如纯图片页），没文字就跳过，不产空 Document
                if text.strip():
                    docs.append(Document(
                        page_content=text,
                        metadata={
                            "source": rel_source,        # 来源路径（引用溯源 + 权限过滤）
                            "parent_id": parent_id,       # 父文档 ID（父子检索回溯用）
                            "page": page_num,             # 页码（答案可标注"见第3页"，增强可信度）
                            "doc_type": "pdf",            # 文档类型（便于按类型筛选/统计）
                        }
                    ))
        return docs

    @staticmethod
    def _ocr_page(page, page_num: int, file_path, lang: str = "chi_sim+eng") -> str:
        """
        OCR 兜底：PyMuPDF 无法提取文字时（扫描件），用 pytesseract 识别。
        依赖 pytesseract + Pillow（可选，未安装则跳过并警告）。

        【为什么是"优雅降级"而不是"强依赖"？】
        OCR 依赖很重（tesseract 引擎 + 语言包几十 MB），且不是所有公司都需要。
        把它设为可选依赖：装了就用，没装就跳过这页，不阻塞主流程。
        生产里扫描件占比通常很低（<5%），牺牲这部分换取部署便利是值得的。
        """
        # 延迟导入 + try：依赖没装时给出明确提示，而不是启动就崩
        try:
            import pytesseract       # OCR 引擎的 Python 封装
            from PIL import Image    # 图像处理（把 pixmap 转成 PIL.Image 才能喂给 tesseract）
        except ImportError:
            # debug 级别而非 warning：避免每次扫描页都刷屏（一个 100 页扫描件会刷 100 次）
            logger.debug(f"OCR 依赖未安装（pytesseract/Pillow），跳过扫描页识别: "
                        f"{getattr(file_path, 'name', file_path)} p{page_num}")
            return ""
        try:
            # dpi=300：渲染分辨率，300 是 OCR 最佳精度（72 太糊识别率低，600 太慢）
            pix = page.get_pixmap(dpi=300)
            # PyMuPDF 的 pixmap 原始字节 → PIL.Image（pytesseract 只认 PIL.Image 或文件路径）
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            # chi_sim+eng：同时识别简体中文和英文（技术文档常含英文术语如"CPU""GPU"）
            text = pytesseract.image_to_string(img, lang=lang)
            if text.strip():
                logger.info(f"  📷 OCR 识别成功: {getattr(file_path, 'name', file_path)} "
                           f"第{page_num}页 ({len(text)} 字符)")
            return text
        except Exception as e:
            # OCR 单页失败不影响其他页（比如某页是纯图案，识别报错）
            logger.warning(f"OCR 失败: {getattr(file_path, 'name', file_path)} "
                          f"第{page_num}页: {e}")
            return ""

    def _load_docx(self, file_path: Path) -> List[Document]:
        """解析 DOCX（Word 文档）

        【局限说明】这里只提取段落文字，不处理 Word 表格/图片。
        完整方案需要额外解析 docx.tables 和 inline shapes，本项目企业文档以段落为主，够用。
        """
        # 别名导入：python-docx 库的类叫 Document，和 LangChain 的 Document 同名，这里改名避免冲突
        from docx import Document as DocxDocument
        docx_doc = DocxDocument(str(file_path))
        # 遍历所有段落，过滤空段落（p.text.strip() 为空的不保留，避免大量空行污染文本）
        # "\n".join：段落之间用换行连接（保留段落结构，便于后续切分）
        text = "\n".join(p.text for p in docx_doc.paragraphs if p.text.strip())
        if text:
            return [Document(
                page_content=text,
                metadata={
                    "source": self._relative_source(file_path),
                    "parent_id": self._parent_id(file_path),
                    "page": 1,          # DOCX 没有稳定页码概念（流式排版），统一记为第1页
                    "doc_type": "docx",
                }
            )]
        return []

    def _load_xlsx(self, file_path: Path) -> List[Document]:
        """
        解析 Excel (.xlsx/.xls)，每行转为自然语言文本 —— 生产级加固版。

        【为什么 Excel 是 RAG 数据接入最难的格式？】
        - 结构化数据（行列），但向量检索基于文本 → 必须把"行"转成"产品名: A，规格: B"的自然语言
        - 企业 Excel 千奇百怪：空 Sheet、合并单元格、列名有换行、重复列名、转置表、超大文件…
        - 每一个坑都会让解析崩溃或产出垃圾数据，必须层层加固

        加固点（每一个都对应真实踩坑）：
          1. 空 Sheet 多重检测（全空/全 NaN/无有效列名/清洗后无数据）
          2. 合并单元格 ffill 填充
          3. 列名清洗（换行、空格、重复列名自动去重）
          4. 转置表自动检测（列多行少 → GPU 型号为列转置为行）
          5. 值清洗（换行、超长截断）
          6. 单 Sheet 异常隔离（一个 Sheet 挂不影响其他）
          7. 行数/内容上限（防内存溢出 + 防单个 Document 过大）
          8. 进度日志（大数据量可监控）
        """
        import pandas as pd   # 局部导入：pandas 很重，只在解析 Excel 时才加载
        docs = []
        parent_id = self._parent_id(file_path)
        rel_source = self._relative_source(file_path)

        # 防御性编程：调用方理论上不会传不存在的文件，但加了更稳（成本几乎为零）
        if not file_path.exists():
            logger.warning(f"Excel 文件不存在: {file_path}")
            return []

        # 【加固8】大文件预警：50MB 的 Excel 解析可能要几十秒，提前告知运维
        # st_size 返回字节数，/ (1024*1024) 转 MB
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > 50:
            logger.warning(f"Excel 文件较大 ({file_size_mb:.1f}MB)，解析可能较慢: {file_path.name}")

        try:
            # ExcelFile 先打开拿 Sheet 名列表（轻量），再用 read_excel 逐个读（按需）
            # engine="openpyxl"：指定解析引擎（xlsx 必须 openpyxl，xls 才是 xlrd）
            excel_file = pd.ExcelFile(file_path, engine="openpyxl")
            sheet_names = excel_file.sheet_names
            logger.info(f"  📊 {file_path.name}: {len(sheet_names)} 个 Sheet")

            processed_sheets = 0
            # enumerate(sheet_names, 1)：Sheet 编号从 1 开始（日志显示 "[2/14]" 更直观）
            for sheet_idx, sheet_name in enumerate(sheet_names, 1):
                # 【加固6】单 Sheet 异常隔离：try 范围限定在单个 Sheet 内
                # 一个 Sheet 崩了（比如格式损坏），不影响其他 13 个 Sheet 继续解析
                try:
                    # ---- 【加固1前置】读取 Sheet ----
                    # dtype=str：所有单元格读成字符串（默认 pandas 会自动推断类型）
                    #   为什么？数字列会被读成 float，"1000000" 变成 "1e+06"（科学计数法）污染文本
                    # header=0：第一行作为列名
                    df = pd.read_excel(
                        file_path, sheet_name=sheet_name,
                        engine="openpyxl", dtype=str, header=0)

                    # ---- 【加固1】空 Sheet 多重检测（企业 Excel 经常有空 Sheet）----
                    # 四道检查层层递进，任一为真就 continue 跳过这个 Sheet：
                    if df.empty or df.isna().all().all():
                        continue                    # ① 完全空或全是 NaN
                    non_empty = df.dropna(how="all")
                    if non_empty.empty:
                        continue                    # ② 删掉全空行后没数据了
                    valid_cols = [c for c in df.columns if str(c).strip()]
                    if not valid_cols:
                        continue                    # ③ 列名全是空白（无效表头）

                    # ---- 【加固3】列名清洗 ----
                    # 企业 Excel 的列名常有这些毛病（真实踩坑）：
                    #   - 换行符："产品\n型号"（Excel 单元格内换行）
                    #   - 前后空格：" 型号 " 和 "型号" 会被当成两列
                    #   - 完全重复：两个列都叫"型号"，pandas 操作时产生歧义
                    clean_headers = []
                    seen = set()                    # 记录已出现的列名，用于去重
                    for col in df.columns:
                        # 统一清洗：strip 去空格 + 替换换行符为空格
                        clean = str(col).strip().replace("\n", " ").replace("\r", " ")
                        if clean in seen:
                            # 重名列自动加后缀：型号 → 型号_2 → 型号_3（避免歧义）
                            suffix = 2
                            while f"{clean}_{suffix}" in seen:
                                suffix += 1
                            clean = f"{clean}_{suffix}"
                        seen.add(clean)
                        clean_headers.append(clean)
                    df.columns = clean_headers      # 用清洗后的列名替换原列名

                    # ---- 【加固2】合并单元格填充 + 空值清洗 ----
                    # ffill（forward fill）：前向填充。Excel 合并单元格只有左上角有值，
                    #   读出来下面都是 NaN，ffill 把上面的值往下填充，恢复完整数据
                    #   例：A1:A3 合并值为"财务部"，读出来是 [财务部, NaN, NaN] → ffill → [财务部, 财务部, 财务部]
                    df.ffill(inplace=True)
                    df = df.fillna("")              # 剩余 NaN（前面没值的）填空串，避免 "nan" 字符串污染文本
                    # 移除全空行（所有列都是空串的行，纯噪音）
                    # axis=1 表示按行操作；lambda 返回 True 的行被 ~ 取反后保留
                    df = df[~df.apply(lambda r: all(v == "" for v in r), axis=1)]
                    if df.empty:
                        continue

                    # ---- 【加固4】转置表检测 ----
                    # 判断条件：列数 > 行数 且 列数 >= 5（启发式，对比表通常列多行少）
                    # 形如：3 行 × 10 列（属性为行、型号为列）→ 转置成 10 行 × 3 列（型号为行）
                    # 详见 _transpose_comparison_table 的注释
                    if df.shape[1] > df.shape[0] and df.shape[1] >= 5:
                        df = _transpose_comparison_table(df)
                        # 转置后可能再次产生 NaN（原表某些属性空），重新填充
                        df.ffill(inplace=True)
                        df = df.fillna("")

                    # ---- 【核心】逐行构建自然语言文本 ----
                    # 【为什么这么转？】向量检索不懂表格，只懂自然语言。
                    # 把一行数据转成 "型号: GB200，架构: Blackwell，显存: 192GB"，
                    # 这样用户问"GB200 显存"，向量能精准命中这一行。
                    columns = list(df.columns)
                    row_docs = 0
                    for idx, row in df.iterrows():  # iterrows：逐行遍历（idx 是行号，row 是该行数据）
                        parts = []
                        for col in columns:
                            val = row[col]
                            if val and str(val).strip():
                                # 【加固5】值清洗：换行符替换 + 超长截断（防单格几万字撑爆 Document）
                                clean_val = str(val).strip().replace("\n", " ").replace("\r", " ")
                                if len(clean_val) > 500:
                                    clean_val = clean_val[:497] + "..."   # 截断并加省略号提示
                                parts.append(f"{col}: {clean_val}")       # "列名: 值" 格式
                        if not parts:
                            continue    # 整行全空，跳过
                        # 行号 idx+2：idx 是 pandas 从 0 开始的行索引，+1 是表头行，+2 才是 Excel 实际行号
                        # （溯源时告诉用户"在第5行"，必须是 Excel 里的真实行号，否则用户找不到）
                        text = f"[{sheet_name}] 第{idx + 2}行: " + "，".join(parts)
                        docs.append(Document(
                            page_content=text,
                            metadata={
                                "source": rel_source,
                                "parent_id": parent_id,
                                "sheet": sheet_name,    # Sheet 名（溯源用）
                                "row": idx + 2,         # Excel 行号（溯源用）
                                "doc_type": "xlsx",
                            }
                        ))
                        row_docs += 1
                        # 【加固7】单 Sheet 行数上限：防恶意/误传的超大表撑爆内存
                        # 10000 行 × 多列 × 多 Sheet 会产生海量 Document，embedding 时长和内存都扛不住
                        if row_docs >= 10000:
                            logger.warning(f"    [{sheet_name}] 超过 10000 行，截断")
                            break

                    # ---- 【关键优化】生成列摘要块 ----
                    # 【为什么需要？】上面的"逐行文本"让每行单独成一个 chunk，
                    # 但用户搜"Blackwell 架构的 GPU 有哪些"时，答案分散在多行，单行 chunk 召回不全。
                    # 列摘要把"架构"这一列的所有值（Blackwell, Hopper, Ampere...）聚成一个 chunk，
                    # 让 Excel 数据在向量检索中能和 PDF 平等竞争（详见 _make_column_summary）
                    summary_doc = self._make_column_summary(
                        df, sheet_name, rel_source, parent_id)
                    if summary_doc:
                        docs.append(summary_doc)

                    processed_sheets += 1
                    # 【加固8】进度日志：14 个 Sheet 逐个打印进度，大数据量时能监控到卡在哪
                    logger.info(f"    [{sheet_idx}/{len(sheet_names)}] {sheet_name}: {row_docs} 行 + 列摘要")

                except Exception as sheet_error:
                    # 【加固6】单 Sheet 异常隔离的核心：异常捕获范围精确到单个 Sheet
                    # 效果：第 5 个 Sheet 崩了，第 6-14 个 Sheet 照常处理
                    logger.warning(f"    ❌ Sheet '{sheet_name}' 解析失败: {sheet_error}")
                    continue

            excel_file.close()    # 显式关闭（虽然 with 更好，但 ExcelFile 这里手动管理）

            # 汇总日志：成功处理了多少 Sheet、产出了多少条目
            if docs:
                unique_sheets = len(set(d.metadata["sheet"] for d in docs))
                logger.info(f"  ✅ {file_path.name}: {processed_sheets}/{len(sheet_names)} Sheet"
                            f" → {len(docs)} 个条目 ({unique_sheets} 个有效 Sheet)")
            else:
                logger.warning(f"  ⚠️ {file_path.name}: 无有效数据")

        except Exception as e:
            # 外层兜底：整个 Excel 文件级别失败（比如文件损坏不是合法 xlsx）
            logger.error(f"解析 Excel 失败 {file_path}: {e}")
            return []
        return docs

    @staticmethod
    def _make_column_summary(df, sheet_name, rel_source, parent_id):
        """
        为 Excel Sheet 生成列摘要块：关键列完整列举，普通列只统计。

        【设计思想：为向量检索补一条"索引项"】
        普通 Excel 行是"一个产品的详情"，但用户常问"有哪些型号/架构"这类聚合性问题。
        如果没有摘要块，答案分散在几百行 chunk 里，向量检索只能召回其中一两行，答不全。
        摘要块把"型号"列的所有值聚成一句话："型号: GB200, B200, H100, ..."，
        用户一搜就命中，相当于给这张表建了个倒排索引。

        效果：用户搜 "Blackwell" 时，向量检索直接命中这个摘要块，
        与 PDF 文本并列竞争，Reranker（或 MMR）会优先选中这个精准块。

        关键列匹配规则：列名含 "型号" / "架构" / "处理器" / "CPU" / "GPU"。
        """
        key_columns = []
        # 识别"关键列"：列名含技术关键词的列才值得做摘要（如"型号""架构"）
        # 其他列（如"备注""序号"）做摘要没意义，反而干扰
        for col in df.columns:
            col_lower = str(col).lower()   # 小写化，匹配更宽松（兼容 "CPU" / "cpu" / "Cpu"）
            if any(kw in col_lower for kw in ("型号", "架构", "处理器", "cpu", "gpu", "model", "产品")):
                key_columns.append(col)

        if not key_columns:
            return None    # 没有关键列的 Sheet 不产摘要块（避免无意义 Document）

        parts = [f"工作表 [{sheet_name}] 关键信息列摘要："]
        for col in key_columns:
            # dropna()：去掉空值；tolist()：转成 Python list
            # 额外过滤掉 "nan" 字符串（dtype=str 读空单元格会变成 "nan"）
            values = [str(v).strip() for v in df[col].dropna().tolist()
                      if str(v).strip() and str(v).strip().lower() != "nan"]
            if not values:
                continue
            # dict.fromkeys(values)：利用 dict 键唯一性去重，且 Python 3.7+ 保持插入顺序
            # 比 set(values) 好：set 无序，去重后型号顺序乱了，可读性差
            unique_vals = list(dict.fromkeys(values))
            # 截断保护：超长摘要会超出 embedding 的 max_seq_length（bge 是 512 token）
            # 100 个值一般够用，更多就截断 + 提示总数
            if len(unique_vals) > 100:
                unique_vals = unique_vals[:80]
                parts.append(f"  {col}: {', '.join(unique_vals)} ... (共 {len(values)} 个值)")
            else:
                parts.append(f"  {col}: {', '.join(unique_vals)}")
        return Document(
            page_content="\n".join(parts),
            metadata={
                "source": rel_source,
                "parent_id": parent_id,
                "sheet": sheet_name,
                "block_type": "column_summary",   # 标记为"摘要块"，检索时可特殊处理
                "doc_type": "xlsx",
            }
        )

    def _load_text(self, file_path: Path, file_type: str) -> List[Document]:
        """解析 MD/TXT（最简单的格式，直接读全文）

        【为什么 MD/TXT 不分页？】
        纯文本没有稳定页码概念，一整个文件就是一个 Document，page=1。
        后续切分由 chunk_documents 的递归切分器按字符长度处理。
        """
        # encoding="utf-8"：显式指定编码（Windows 默认 GBK，不指定会乱码）
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        return [Document(
            page_content=content,
            metadata={
                "source": self._relative_source(file_path),
                "parent_id": self._parent_id(file_path),
                "page": 1,
                "doc_type": file_type,   # "markdown" 或 "text"，便于按格式筛选
            }
        )]

    def _relative_source(self, file_path: Path) -> str:
        r"""相对 data_path 的源路径（保留部门目录，供权限过滤）

        【为什么用相对路径而不是绝对路径？】
        1. 跨环境一致：开发机路径 F:\code\... 到生产机变成 /app/...，绝对路径会变
        2. 保留部门信息：相对路径是 "财务/报销.pdf"，第一段就是部门，权限过滤要用
        绝对路径可能是 "F:\code\data\docs\财务\报销.pdf"，部门在中间，提取麻烦
        """
        try:
            # relative_to：算出相对于 _data_root 的相对路径
            # as_posix()：统一用 / 分隔符（Windows 默认 \，跨平台不一致）
            return file_path.resolve().relative_to(self._data_root).as_posix()
        except Exception:
            # 兜底：如果算不出相对路径（比如文件不在 data_root 下），退化成绝对路径
            return file_path.as_posix()

    def _parent_id(self, file_path: Path) -> str:
        """生成确定性 parent_id（基于相对路径，供父子检索）

        【为什么用 md5 而不是 uuid？】
        parent_id 要"确定性"：同一个文件每次构建都得到相同的 parent_id，
        这样增量索引才能对比"这个 chunk 属于哪个父文档"。
        uuid 每次随机，无法对比。md5(相对路径) 保证：同路径→同ID。
        """
        try:
            relative = file_path.resolve().relative_to(self._data_root).as_posix()
        except Exception:
            relative = file_path.as_posix()
        # md5 + hexdigest：32 位十六进制字符串，作为稳定的唯一标识
        return hashlib.md5(relative.encode("utf-8")).hexdigest()

    def _enhance_metadata(self, doc: Document):
        """
        增强企业元数据：从路径提取部门/分类
        约定：data/docs/{部门}/{文件}  →  department=部门
        对应《真实RAG全貌》⑤权限（RBAC 过滤用 department）

        【这是整个权限系统的数据基础！】
        RBAC 部门过滤（retrieval_optimization.py）就是读 department 元数据决定能看哪些文档。
        如果这里没正确写入 department，权限隔离形同虚设。
        """
        source = doc.metadata.get("source", "")
        parts = Path(source).parts
        # 路径形如 "HR制度.pdf"（无部门，在根目录）或 "财务/报销.pdf"（财务部门）
        # parts[-2]：倒数第二段就是部门目录名（最后一段是文件名）
        if len(parts) > 1:
            doc.metadata["department"] = parts[-2]
        else:
            doc.metadata["department"] = "公共"   # 没有部门目录的算"公共"，所有人可见

    # ===== ② 文档处理：切分（表格感知 + 父子分块）=====
    def _split_text_and_tables(self, content: str) -> List[dict]:
        """
        表格感知预切分：把文档拆成「文本块」和「表格块」序列。
        表格（连续的 |...| 行）作为一个整体块，绝不切断。

        【为什么需要这个？—— 对应踩坑实录坑22】
        标准 RAG 用 RecursiveCharacterTextSplitter 按字符长度切分，它不认识表格结构。
        一个 6 行的表格（每行一个型号参数），如果遇到 chunk_size 边界，会被从中间切断：
            切点在第 3 行 → 前 3 行进 chunk A，后 3 行进 chunk B
            用户问"P2 的参数" → P2 在第 3 行恰好在 chunk A，但它的完整数据在第 3-6 行
        实测表格完整召回率从 50% 提升到 100%。

        【算法思路：单遍扫描 + 双指针】
        遍历每一行，遇到表格行（以 | 开头）就把连续的表格行整体收集，
        遇到文本行就攒到缓冲区。最后产出交替的 text/table 块序列。

        返回: [{"type": "text"|"table", "content": "..."}]
        """
        blocks = []
        lines = content.split("\n")
        i = 0                    # 主指针：当前扫描到哪一行
        n = len(lines)
        text_buf = []            # 文本缓冲区：积累连续的非表格行，攒够一段后批量产出

        def flush_text():
            """把缓冲区的文本行打包成一个 text 块，然后清空缓冲区。

            闭包捕获 text_buf：这是 Python 闭包修改外部变量的标准模式
            （直接读没问题，但要修改 list 内容用 .clear() 而非 = []）"""
            if text_buf:
                blocks.append({"type": "text", "content": "\n".join(text_buf).strip()})
                text_buf.clear()

        # 单遍扫描：i 从 0 走到 n-1，每个分支都会推进 i，不会死循环
        while i < n:
            line = lines[i]
            # 识别表格行：Markdown 表格的特征是"以 | 开头"（如 "| 型号 | 显存 |"）
            # 加 "and" 条件：必须同时含 | 才算（避免误判，虽然 startswith("|") 已基本够）
            if "|" in line and line.strip().startswith("|"):
                # 内层循环：向后贪心收集所有连续的表格行（j 是表格块的结束指针）
                # 整个表格块作为一个整体，绝不切断
                table_lines = [line]
                j = i + 1
                while j < n and lines[j].strip().startswith("|"):
                    table_lines.append(lines[j])
                    j += 1
                # 关键：产出表格块前，先把之前攒的文本 flush（保证 text/table 顺序正确）
                flush_text()
                blocks.append({"type": "table", "content": "\n".join(table_lines)})
                i = j            # 主指针跳过整个表格块，继续扫描后面的内容
            else:
                # 非表格行：攒到缓冲区，等遇到表格或文档结束时再批量产出
                text_buf.append(line)
                i += 1
        # 循环结束：把最后一段还没 flush 的文本产出（尾部兜底，否则会丢内容）
        flush_text()
        return blocks

    def chunk_documents(self) -> List[Document]:
        """
        表格感知切分（文本递归切 + 表格整体保留），保留父子关系
        对应《真实RAG全貌》②结构感知切分 + 坑22（表格不切断）

        【整体流程】
        for 每个父文档 doc:
            1. _split_text_and_tables 拆成 text/table 块序列
            2. text 块 → 递归切分器切成小 chunk；table 块 → 整体一个 chunk
            3. 给每个 chunk 打元数据（chunk_id / parent_id / content_hash）
        """
        logger.info(f"正在进行表格感知切分（size={self.chunk_size}, overlap={self.chunk_overlap}）...")
        if not self.documents:
            raise ValueError("请先加载文档")

        # 创建递归字符切分器（一次创建，复用，避免每个 doc 都重建对象）
        # separators 是分隔符优先级列表（从前往后尝试）：
        #   "\n\n"：段落边界（最优先，切出来语义最完整）
        #   "\n"：行边界
        #   "。"：中文句号（中文文档主切分点）
        #   "；"：分号（子句边界）
        #   "，"：逗号（更细的子句）
        #   " "：空格（英文/数字边界）
        #   ""：单字符（最后兜底，保证一定能切到 chunk_size 以内）
        # 【设计哲学】优先在语义边界切，实在不行才硬切，保证 chunk 尽量完整
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", "。", "；", "，", " ", ""],
        )

        all_chunks = []
        for doc in self.documents:
            try:
                # 取父文档 ID（_load_xxx 时已写入；没有则临时生成，但这种情况不该出现）
                parent_id = doc.metadata.get("parent_id", str(uuid.uuid4()))
                # ① 表格感知预切分：拆成 [{type, content}, ...]
                blocks = self._split_text_and_tables(doc.page_content)
                doc_chunks = []
                # 复制父文档的元数据作为子块的基础元数据（source/department/page 等都要继承）
                # dict() 浅拷贝：避免修改子块元数据时污染父文档
                base_meta = dict(doc.metadata)
                # ② 对每个 block 分别处理
                for block in blocks:
                    if block["type"] == "table":
                        # 表格块：整体作为一个 chunk，不递归切分（坑22 的核心修复）
                        # {**base_meta, "block_type": "table"}：合并父元数据 + 标记为表格类型
                        doc_chunks.append(Document(page_content=block["content"],
                                              metadata={**base_meta, "block_type": "table"}))
                    else:
                        # 文本块：递归切分成多个小 chunk
                        # split_text 返回字符串列表，每个字符串是一个切好的 chunk
                        sub = text_splitter.split_text(block["content"])
                        for s in sub:
                            doc_chunks.append(Document(page_content=s,
                                                  metadata={**base_meta, "block_type": "text"}))

                # ③ 给每个子块补充 chunk 级元数据（索引构建和父子检索都要用）
                for i, chunk in enumerate(doc_chunks):
                    child_id = str(uuid.uuid4())   # 子块唯一 ID（随机，不要求确定性）
                    chunk.metadata.update({
                        "chunk_id": child_id,
                        "parent_id": parent_id,     # 指回父文档（父子检索回溯用）
                        "chunk_index": i,           # 在父文档内的序号（溯源：第几个块）
                        "chunk_size": len(chunk.page_content),   # 字符长度（统计/异常检测用）
                        # 【关键】content_hash：增量索引用的核心字段
                        # index_construction.build_incremental 对比新旧 hash，相同则跳过 embedding
                        "content_hash": hashlib.md5(chunk.page_content.encode("utf-8")).hexdigest(),
                    })
                    self.parent_child_map[child_id] = parent_id   # 建立映射表
                all_chunks.extend(doc_chunks)
            except Exception as e:
                # 单文档切分失败：降级把整篇文档作为一个 chunk（保证不丢内容）
                logger.warning(f"切分失败 {doc.metadata.get('source')}: {e}")
                all_chunks.append(doc)

        self.chunks = all_chunks
        logger.info(f"切分完成，共 {len(all_chunks)} 个 chunk")
        return all_chunks

    def get_parent_documents(self, child_chunks: List[Document]) -> List[Document]:
        """
        根据子块回溯父文档（去重、按相关性排序）
        父子检索亮点：小块精准检索 → 父块完整上下文

        【为什么要父子检索？】
        - 小 chunk（500 字）向量检索精准，但上下文不足，LLM 答题可能缺关键信息
        - 大 chunk（整篇文档）上下文足，但向量稀释，检索召回率低
        - 折中：用小 chunk 检索（保证召回），命中后回溯父文档拿完整上下文给 LLM
        - 这就是业界说的 "Small-to-Big" / "Parent-Child Retrieval" 模式

        【相关性排序的巧思】
        如果一个父文档被多个子块命中（比如 3 个子块都指向它），说明这个父文档整体高度相关。
        按命中次数排序，最相关的父文档排前面，优先喂给 LLM。
        """
        # parent_relevance：parent_id → 命中次数（被几个子块指向）
        parent_relevance: Dict[str, int] = {}
        # parent_docs_map：parent_id → 父文档对象（去重，一个父文档只存一份）
        parent_docs_map: Dict[str, Document] = {}

        for chunk in child_chunks:
            parent_id = chunk.metadata.get("parent_id")
            if parent_id:
                # 命中计数 +1（get 默认 0，+1 后写回）
                parent_relevance[parent_id] = parent_relevance.get(parent_id, 0) + 1
                # 第一次见到这个 parent_id，去 documents 里找出对应的父文档缓存起来
                # （避免同一个 parent_id 多次遍历 documents 列表，O(n²) 退化成 O(n)）
                if parent_id not in parent_docs_map:
                    for doc in self.documents:
                        if doc.metadata.get("parent_id") == parent_id:
                            parent_docs_map[parent_id] = doc
                            break

        # 按命中次数降序排序（最相关的父文档排前面）
        sorted_ids = sorted(parent_relevance.keys(), key=lambda x: parent_relevance[x], reverse=True)
        # 从 map 里取出父文档（按排序后的顺序），过滤掉找不到的（理论上不会，防御性编程）
        parent_docs = [parent_docs_map[pid] for pid in sorted_ids if pid in parent_docs_map]

        logger.info(f"从 {len(child_chunks)} 子块回溯到 {len(parent_docs)} 个父文档")
        return parent_docs

    def get_statistics(self) -> Dict[str, Any]:
        """统计信息（监控/日志用）

        【运维价值】构建完知识库后打印这份统计，能快速发现异常：
        - total_documents / total_chunks 比例不对？可能切分参数有问题
        - departments 分布异常？某部门文档没加载进来
        - avg_chunk_size 过大/过小？chunk_size 配置需要调整
        """
        # 按部门统计文档数（验证 RBAC 数据基础是否完整）
        departments = {}
        for doc in self.documents:
            dept = doc.metadata.get("department", "未知")
            departments[dept] = departments.get(dept, 0) + 1
        return {
            "total_documents": len(self.documents),
            "total_chunks": len(self.chunks),
            "departments": departments,
            # 平均 chunk 大小：判断切分质量的核心指标
            # 过大（>1000）：可能很多表格没被切分（正常，表格整体保留）
            # 过小（<100）：切得太碎，上下文割裂
            "avg_chunk_size": (
                sum(c.metadata.get("chunk_size", 0) for c in self.chunks) / len(self.chunks)
                if self.chunks else 0    # chunks 为空时返回 0，避免除零错误
            ),
        }
