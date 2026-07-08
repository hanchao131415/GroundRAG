"""生成 GroundRAG 合成样例知识库，覆盖 md/txt/xlsx/docx/pdf 五种格式。

运行：python scripts/make_sample_docs.py
依赖：pandas, openpyxl, python-docx, PyMuPDF(fitz) —— 均为后端已有依赖。
PDF 用 fitz 内置 CJK 字体 china-s，无需额外字体文件。
"""
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent / "data" / "docs"


def write_text(rel, content):
    p = ROOT / rel
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(content, encoding="utf-8")


def write_xlsx(rel, header, rows, sheet="Sheet1"):
    import pandas as pd
    p = ROOT / rel
    p.parent.mkdir(parents=True, exist_ok=True)
    pd.DataFrame(rows, columns=header).to_excel(p, sheet_name=sheet, index=False)


def write_docx(rel, title, paragraphs):
    from docx import Document
    p = ROOT / rel
    p.parent.mkdir(parents=True, exist_ok=True)
    d = Document()
    d.add_heading(title, level=1)
    for para in paragraphs:
        d.add_paragraph(para)
    d.save(str(p))


def write_pdf(rel, title, lines):
    import fitz  # PyMuPDF，内置 CJK 字体 china-s
    p = ROOT / rel
    p.parent.mkdir(parents=True, exist_ok=True)
    doc = fitz.open()
    page = doc.new_page()
    y = 60
    page.insert_text((72, y), title, fontname="china-s", fontsize=18)
    y += 30
    for ln in lines:
        page.insert_text((72, y), ln, fontname="china-s", fontsize=12)
        y += 20
    doc.save(str(p))
    doc.close()


if __name__ == "__main__":
    # ===== HR：md + txt + pdf =====
    write_text("HR/年假管理制度.md",
               "# 示例科技有限公司 年假管理制度\n\n## 年假天数标准\n"
               "- 工作满1年不满3年：年假5天\n- 工作满3年不满5年：年假10天\n"
               "- 工作满5年不满10年：年假15天\n- 工作满10年以上：年假20天\n\n"
               "## 请假流程\n请假须提前通过OA系统提交申请，由直属主管审批。年假可按半天为单位使用。\n")
    write_text("HR/考勤管理制度.txt",
               "示例科技有限公司 考勤管理制度\n\n"
               "工作时间：9:00-18:00，午休 12:00-13:00。\n"
               "打卡规则：每日上班、下班各打卡一次。迟到超过30分钟按半天事假处理。病假须提供医院证明。\n")
    write_pdf("HR/薪酬福利管理办法.pdf", "薪酬福利管理办法",
              ["工资每月10日发放，遇节假日顺延。",
               "年终奖根据公司业绩与个人考核发放。",
               "五险一金按国家规定缴纳。"])

    # ===== 财务：xlsx + docx =====
    write_xlsx("财务/差旅报销标准.xlsx",
               header=["城市级别", "住宿费上限(元/晚)", "市内交通补贴(元/天)"],
               rows=[["一线城市", 500, 80], ["二线城市", 400, 60], ["三线及以下", 300, 50]])
    write_docx("财务/报销管理制度.docx", "报销管理制度",
               ["所有报销须在费用发生后30日内提交。",
                "单笔超过5000元的报销须部门负责人与财务总监双签。",
                "发票须真实、合法，与业务一致。"])

    # ===== IT：pdf + md =====
    write_pdf("IT/信息安全管理规定.pdf", "信息安全管理规定",
              ["密码长度至少8位，须包含大小写字母与数字。",
               "密码每90天更换一次。",
               "连续5次输错密码将锁定账号30分钟。",
               "禁止将公司内网账号借予他人。离职时须立即收回所有系统权限。"])
    write_text("IT/网络管理规范.md",
               "# 示例科技有限公司 网络管理规范\n\n"
               "## 办公网络\n办公网络与访客网络物理隔离。禁止私自接入路由器或随身WiFi。\n\n"
               "## VPN\n远程办公须通过公司VPN接入，VPN账号禁止共享。\n")

    # ===== 行政：docx + xlsx =====
    write_docx("行政/办公用品领用管理办法.docx", "办公用品领用管理办法",
               ["办公用品统一由行政部采购。员工通过OA系统提交领用申请，主管审批后到行政部领取。",
                "每人每月办公用品预算不超过200元。特殊用品（如人体工学椅）须部门负责人额外审批。"])
    write_xlsx("行政/采购分级标准.xlsx",
               header=["采购金额(元)", "审批层级", "是否需招标"],
               rows=[["<1000", "部门主管", "否"],
                     ["1000-10000", "部门负责人", "否"],
                     [">10000", "总经理", "是"]])

    # ===== 研发：md + txt =====
    write_text("研发/代码规范.md",
               "# 示例科技有限公司 研发代码规范\n\n## 代码评审\n"
               "所有代码合并前须经至少一位同事评审。主分支禁止直接 push。\n\n"
               "## 分支命名\n功能分支：feature/简短描述。修复分支：fix/简短描述。发布分支：release/版本号。\n")
    write_text("研发/开发流程.txt",
               "示例科技有限公司 研发开发流程\n\n"
               "需求评审 → 技术方案 → 编码 → 代码评审 → 测试 → 灰度发布 → 全量发布。\n"
               "每个环节须有产出物并存档。线上发布须双人复核。\n")

    print(f"sample docs written under {ROOT}")
